"""
Document Parsers
Extracts text from various document formats
"""
import os
from typing import Dict, Any
from pathlib import Path
import PyPDF2
from docx import Document
import markdown
from bs4 import BeautifulSoup


class DocumentParser:
    """Base document parser"""

    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        """
        Parse document and extract text

        Args:
            file_path: Path to document file

        Returns:
            Dictionary with extracted text and metadata
        """
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            return PDFParser.parse(file_path)
        elif ext in [".docx", ".doc"]:
            return DOCXParser.parse(file_path)
        elif ext == ".txt":
            return TXTParser.parse(file_path)
        elif ext == ".md":
            return MarkdownParser.parse(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")


class PDFParser:
    """PDF document parser"""

    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        """Extract text from PDF"""
        text_content = []
        metadata = {
            "num_pages": 0,
            "title": None,
            "author": None,
        }

        try:
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                metadata["num_pages"] = len(reader.pages)

                # Extract metadata
                if reader.metadata:
                    metadata["title"] = reader.metadata.get("/Title")
                    metadata["author"] = reader.metadata.get("/Author")

                # Extract text from all pages
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

            full_text = "\n\n".join(text_content)

            return {
                "text": full_text,
                "metadata": metadata,
                "format": "pdf",
                "success": True,
                "error": None
            }

        except Exception as e:
            return {
                "text": "",
                "metadata": metadata,
                "format": "pdf",
                "success": False,
                "error": str(e)
            }


class DOCXParser:
    """DOCX document parser"""

    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX"""
        try:
            doc = Document(file_path)

            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        tables_text.append(row_text)

            # Combine all text
            full_text = "\n\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n" + "\n".join(tables_text)

            metadata = {
                "num_paragraphs": len(paragraphs),
                "num_tables": len(doc.tables),
            }

            return {
                "text": full_text,
                "metadata": metadata,
                "format": "docx",
                "success": True,
                "error": None
            }

        except Exception as e:
            return {
                "text": "",
                "metadata": {},
                "format": "docx",
                "success": False,
                "error": str(e)
            }


class TXTParser:
    """Plain text parser"""

    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        """Extract text from TXT"""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                text = file.read()

            # Try different encodings if UTF-8 fails
            if not text:
                with open(file_path, "r", encoding="latin-1") as file:
                    text = file.read()

            lines = text.split("\n")
            metadata = {
                "num_lines": len(lines),
                "encoding": "utf-8"
            }

            return {
                "text": text,
                "metadata": metadata,
                "format": "txt",
                "success": True,
                "error": None
            }

        except Exception as e:
            return {
                "text": "",
                "metadata": {},
                "format": "txt",
                "success": False,
                "error": str(e)
            }


class MarkdownParser:
    """Markdown parser"""

    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        """Extract text from Markdown"""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                md_content = file.read()

            # Convert markdown to HTML
            html = markdown.markdown(md_content)

            # Extract plain text from HTML
            soup = BeautifulSoup(html, "html.parser")
            plain_text = soup.get_text()

            # Also keep the original markdown
            metadata = {
                "has_html": bool(html),
                "markdown_length": len(md_content),
            }

            return {
                "text": plain_text,
                "markdown": md_content,
                "html": html,
                "metadata": metadata,
                "format": "markdown",
                "success": True,
                "error": None
            }

        except Exception as e:
            return {
                "text": "",
                "metadata": {},
                "format": "markdown",
                "success": False,
                "error": str(e)
            }
